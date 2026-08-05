"""
DOCX, Markdown, and plain-text extraction
=========================================
Produces the same `ParsedDocument` shape as the PDF extractor so chunking
never needs to know the source format.

The previous DOCX loader read only ``doc.paragraphs``, which silently dropped
every table in the file. This one walks the document body in order, so tables
appear as markdown blocks in their true position relative to the prose.
"""

from __future__ import annotations

import logging
import re
import time
from pathlib import Path

from .models import Block, BlockKind, ExtractionStats, ParsedDocument, file_content_hash
from .normalize import normalize_block_text

logger = logging.getLogger(__name__)

EXTRACTOR_VERSION = "office/1.0"

_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_MD_FENCE = re.compile(r"^\s*```")
_MD_TABLE_ROW = re.compile(r"^\s*\|.*\|\s*$")
_LIST_START = re.compile(r"^\s*(?:[-*+]|\d+[.)])\s+")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _docx_table_to_markdown(table) -> str:
    """Render a python-docx table as markdown."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
        if any(cells):
            rows.append(cells)
    if len(rows) < 2:
        return ""

    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header, *body = rows
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    lines.extend("| " + " | ".join(r) + " |" for r in body)
    return "\n".join(lines)


def _docx_heading_level(style_name: str) -> int | None:
    """Map a Word paragraph style to a heading level, or None if not a heading."""
    if not style_name:
        return None
    m = re.match(r"^Heading (\d)$", style_name.strip())
    if m:
        return int(m.group(1))
    if style_name.strip() in ("Title", "Subtitle"):
        return 1
    return None


def extract_docx(path: str) -> ParsedDocument:
    """Extract a .docx into blocks, preserving paragraph/table document order."""
    started = time.perf_counter()
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RuntimeError("python-docx is required to read .docx files") from exc

    doc = Document(path)
    stats = ExtractionStats(page_count=1)
    blocks: list[Block] = []

    # Walk the body's XML children so paragraphs and tables stay interleaved.
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]

        if tag == "p":
            para = Paragraph(child, doc)
            raw = para.text
            if not raw.strip():
                continue
            clean, counters = normalize_block_text(raw)
            stats.dehyphenated += counters["dehyphenated"]
            stats.dropped_noise_lines += counters["dropped_noise_lines"]
            if not clean.strip():
                continue

            style = getattr(getattr(para, "style", None), "name", "") or ""
            level = _docx_heading_level(style)
            if level is not None:
                kind = BlockKind.HEADING
                stats.heading_count += 1
            elif _LIST_START.match(clean) or style.startswith("List"):
                kind, level = BlockKind.LIST_ITEM, None
            else:
                kind, level = BlockKind.PARAGRAPH, None

            blocks.append(Block(text=clean, kind=kind, page=1, level=level))

        elif tag == "tbl":
            md = _docx_table_to_markdown(Table(child, doc))
            if md:
                stats.table_count += 1
                blocks.append(Block(text=md, kind=BlockKind.TABLE, page=1))

    stats.block_count = len(blocks)
    stats.duration_ms = int((time.perf_counter() - started) * 1000)

    return ParsedDocument(
        source_path=str(path),
        filename=Path(path).name,
        title=_first_heading(blocks) or Path(path).stem,
        content_hash=file_content_hash(path),
        blocks=blocks,
        stats=stats,
        extractor=EXTRACTOR_VERSION,
    )


# ---------------------------------------------------------------------------
# Markdown / plain text
# ---------------------------------------------------------------------------


def extract_markdown(path: str) -> ParsedDocument:
    """
    Extract .md or .txt.

    Markdown ATX headings become HEADING blocks, fenced code becomes CODE, and
    pipe tables are kept intact as TABLE blocks. Plain .txt files simply have
    none of those patterns and fall through as paragraphs.
    """
    started = time.perf_counter()
    raw = Path(path).read_text(encoding="utf-8", errors="ignore")

    stats = ExtractionStats(page_count=1)
    blocks: list[Block] = []
    buffer: list[str] = []
    buffer_kind = BlockKind.PARAGRAPH

    def flush() -> None:
        nonlocal buffer, buffer_kind
        if not buffer:
            return
        text = "\n".join(buffer).strip()
        buffer = []
        kind, buffer_kind = buffer_kind, BlockKind.PARAGRAPH
        if not text:
            return
        if kind is BlockKind.CODE:
            blocks.append(Block(text=text, kind=kind, page=1))
            return
        clean, counters = normalize_block_text(text)
        stats.dehyphenated += counters["dehyphenated"]
        stats.dropped_noise_lines += counters["dropped_noise_lines"]
        if clean.strip():
            if kind is BlockKind.TABLE:
                stats.table_count += 1
            blocks.append(Block(text=clean, kind=kind, page=1))

    in_fence = False
    for line in raw.split("\n"):
        if _MD_FENCE.match(line):
            if in_fence:
                buffer.append(line)
                flush()
                in_fence = False
            else:
                flush()
                in_fence = True
                buffer_kind = BlockKind.CODE
                buffer.append(line)
            continue

        if in_fence:
            buffer.append(line)
            continue

        heading = _MD_HEADING.match(line)
        if heading:
            flush()
            title = heading.group(2).strip()
            if title:
                stats.heading_count += 1
                blocks.append(
                    Block(text=title, kind=BlockKind.HEADING, page=1, level=len(heading.group(1)))
                )
            continue

        if _MD_TABLE_ROW.match(line):
            if buffer_kind is not BlockKind.TABLE:
                flush()
                buffer_kind = BlockKind.TABLE
            buffer.append(line)
            continue

        if not line.strip():
            flush()
            continue

        if buffer_kind is BlockKind.TABLE:
            flush()
        if _LIST_START.match(line) and not buffer:
            buffer_kind = BlockKind.LIST_ITEM
        buffer.append(line)

    flush()

    stats.block_count = len(blocks)
    stats.duration_ms = int((time.perf_counter() - started) * 1000)

    return ParsedDocument(
        source_path=str(path),
        filename=Path(path).name,
        title=_first_heading(blocks) or Path(path).stem,
        content_hash=file_content_hash(path),
        blocks=blocks,
        stats=stats,
        extractor=EXTRACTOR_VERSION,
    )


def _first_heading(blocks: list[Block]) -> str | None:
    for b in blocks[:10]:
        if b.is_heading:
            return b.text.strip().split("\n")[0][:160]
    return None
